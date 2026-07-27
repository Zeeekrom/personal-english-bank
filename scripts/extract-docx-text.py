"""Extract complete readable text from a DOCX without modifying the source."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


def iter_blocks(parent: DocumentObject):
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract(path: Path) -> dict[str, object]:
    document = Document(path)
    blocks: list[str] = []
    paragraph_count = 0
    table_count = 0

    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            paragraph_count += 1
            text = block.text.strip()
            if text:
                blocks.append(text)
            continue

        table_count += 1
        for row in block.rows:
            cells = [
                " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                for cell in row.cells
            ]
            if any(cells):
                blocks.append("\t".join(cells))

    for section_index, section in enumerate(document.sections, start=1):
        header_text = "\n".join(
            paragraph.text.strip()
            for paragraph in section.header.paragraphs
            if paragraph.text.strip()
        )
        footer_text = "\n".join(
            paragraph.text.strip()
            for paragraph in section.footer.paragraphs
            if paragraph.text.strip()
        )
        if header_text:
            blocks.insert(0, f"[Header {section_index}]\n{header_text}")
        if footer_text:
            blocks.append(f"[Footer {section_index}]\n{footer_text}")

    text = "\n\n".join(blocks).strip()
    return {
        "text": text,
        "characters": len(text),
        "paragraphs": paragraph_count,
        "tables": table_count,
    }


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: extract-docx-text.py <input.docx>")
    sys.stdout.reconfigure(encoding="utf-8")
    result = extract(Path(sys.argv[1]))
    json.dump(result, sys.stdout, ensure_ascii=False)


if __name__ == "__main__":
    main()
